"""
DOCX Exporter (TASK-016).
Exports a ResearchReport as a professionally formatted .docx file.
"""

import asyncio
import logging
from pathlib import Path

from app.models.report import ResearchReport

logger = logging.getLogger(__name__)


class DocxExporterError(Exception):
    """Raised when DOCX export fails."""


class DocxExporter:
    """
    Exports a ResearchReport to a .docx file using python-docx.

    Args:
        output_dir: Directory where the file will be saved.
    """

    def __init__(self, output_dir: str | Path = "./outputs"):
        self.output_dir = Path(output_dir)

    async def export(self, report: ResearchReport) -> Path:
        """
        Export the report to a .docx file.

        Args:
            report: The fully assembled ResearchReport.

        Returns:
            Path to the saved .docx file.

        Raises:
            DocxExporterError: On any file-creation failure.
        """
        try:
            return await asyncio.to_thread(self._build_docx, report)
        except Exception as e:
            raise DocxExporterError(f"Failed to export DOCX: {e}") from e

    def _build_docx(self, report: ResearchReport) -> Path:
        """Synchronous DOCX construction (runs in a thread)."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # -----------------------------------------------------------------
        # Styles
        # -----------------------------------------------------------------
        styles = doc.styles

        # Title
        title_para = doc.add_heading(report.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Metadata line
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(
            f"Research Date: {report.research_date.strftime('%Y-%m-%d')}  |  "
            f"Depth: {report.depth.capitalize()}  |  "
            f"Sources Included: {report.sources_included}"
        ).italic = True

        doc.add_paragraph()  # spacer

        # -----------------------------------------------------------------
        # Executive Summary
        # -----------------------------------------------------------------
        if report.executive_summary:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(report.executive_summary)
            doc.add_paragraph()

        # -----------------------------------------------------------------
        # Sections
        # -----------------------------------------------------------------
        sorted_sections = sorted(report.sections, key=lambda s: s.order)
        for section in sorted_sections:
            doc.add_heading(section.heading, level=1)

            # Body content
            if section.content:
                for para_text in section.content.split("\n"):
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    if para_text.startswith("* "):
                        # Bullet point
                        p = doc.add_paragraph(para_text[2:], style="List Bullet")
                    elif para_text.startswith("## "):
                        doc.add_heading(para_text[3:], level=2)
                    elif para_text.startswith("### "):
                        doc.add_heading(para_text[4:], level=3)
                    else:
                        doc.add_paragraph(para_text)

            # Inline findings (Key Findings section)
            if section.findings:
                for finding in section.findings:
                    if finding.uncertain:
                        p = doc.add_paragraph()
                        p.add_run("⚠ Uncertain: ").bold = True
                        p.add_run(finding.claim)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(finding.claim)

                    # Citation markers bold inline
                    if finding.citations:
                        markers = " ".join(c.marker for c in finding.citations)
                        run = p.add_run(f"  {markers}")
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x00, 0x70, 0xBB)

            doc.add_paragraph()

        # -----------------------------------------------------------------
        # Save
        # -----------------------------------------------------------------
        self.output_dir.mkdir(parents=True, exist_ok=True)
        safe_title = "".join(c for c in report.title if c.isalnum() or c in " _-")[:60].strip()
        filename = f"{safe_title}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        logger.info(f"DOCX saved to: {filepath}")
        return filepath
