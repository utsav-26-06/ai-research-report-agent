"""
Unit tests for DocxExporter (TASK-016).
"""

from __future__ import annotations

import os
import tempfile
import pytest

from app.export.docx_exporter import DocxExporter, DocxExporterError
from app.models.report import Citation, Finding, ReportSection, ResearchReport


def _make_report() -> ResearchReport:
    citation = Citation(
        marker="[1]",
        source_id="src1",
        chunk_id="chk1",
        url="http://example.com",
        title="Example",
        domain="example.com",
    )
    finding = Finding(
        claim="AI is advancing rapidly. [1]",
        evidence="Models are smarter.",
        citations=[citation],
        confidence=0.9,
        sub_question_id="sq-1",
    )
    sections = [
        ReportSection(heading="Introduction", content="This is the intro.", order=0),
        ReportSection(heading="Key Findings", content="* Bullet 1\n* Bullet 2", findings=[finding], order=1),
        ReportSection(heading="Analysis", content="Detailed analysis here.", order=2),
        ReportSection(heading="Conclusion", content="Final thoughts.", order=3),
        ReportSection(heading="References", content="[1] Example. example.com. Retrieved from http://example.com", order=4),
    ]
    return ResearchReport(
        title="AI Research Report",
        research_question="How is AI evolving?",
        depth="standard",
        executive_summary="AI is evolving rapidly.",
        sections=sections,
        all_citations=[citation],
        sources_included=1,
    )


@pytest.mark.asyncio
async def test_docx_file_created():
    """The exporter should produce a non-empty .docx file."""
    with tempfile.TemporaryDirectory() as tmpdir:
        exporter = DocxExporter(output_dir=tmpdir)
        report = _make_report()

        filepath = await exporter.export(report)

        assert filepath.exists()
        assert filepath.suffix == ".docx"
        assert filepath.stat().st_size > 0


@pytest.mark.asyncio
async def test_docx_all_sections_present():
    """All section headings should appear in the DOCX document text."""
    from docx import Document

    with tempfile.TemporaryDirectory() as tmpdir:
        exporter = DocxExporter(output_dir=tmpdir)
        report = _make_report()
        filepath = await exporter.export(report)

        doc = Document(str(filepath))
        all_text = " ".join(p.text for p in doc.paragraphs)

        for section in report.sections:
            assert section.heading in all_text, f"Section '{section.heading}' not in DOCX"


@pytest.mark.asyncio
async def test_docx_heading_hierarchy():
    """Title should be a heading; intro should appear as a heading level 1."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with tempfile.TemporaryDirectory() as tmpdir:
        exporter = DocxExporter(output_dir=tmpdir)
        report = _make_report()
        filepath = await exporter.export(report)

        doc = Document(str(filepath))
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
        assert "AI Research Report" in heading_texts
        assert "Introduction" in heading_texts
        assert "Key Findings" in heading_texts

